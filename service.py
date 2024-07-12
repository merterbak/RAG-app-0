import os
import PyPDF2
from docx import Document
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from typing import List
from llama_index.core import (
    ServiceContext,
    StorageContext,
    VectorStoreIndex,
    load_index_from_storage,
    Document as LlamaDocument
)
from llama_index.core.node_parser import (
    HierarchicalNodeParser,
    get_leaf_nodes,
)
from llama_index.core.retrievers import AutoMergingRetriever
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.llms.openai import OpenAI
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
import io
from dotenv import load_dotenv
import logging
from llama_index.core import Settings, set_global_service_context
app = FastAPI()

load_dotenv()
os.environ["TOKENIZERS_PARALLELISM"] = "false"
logging.basicConfig(level=logging.INFO) 
logger = logging.getLogger(__name__)

class User:
    def __init__(self, username):
        self.username = username
        self.llm = "gpt-3.5-turbo"
        self.embedder = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"

@app.post("/document-uploader/")
async def upload_documents(username: str = Form(...), files: List[UploadFile] = File(...)):
    user = User(username=username)
    documents = []
    for file in files:
        text = await _extract_text_from_document(file)
        documents.append(LlamaDocument(text=text, id_=f"doc_{user.username}_{file.filename}"))
    
    await _create_or_update_index(user, documents)
    return {"message": "Documents are uploaded and indexed successfully."}

async def _extract_text_from_document(file: UploadFile) -> str:
    byte_object = await file.read()
    file_name = file.filename
    file_extension = os.path.splitext(file_name)[1].lower()
    if file_extension == '.txt':
        return byte_object.decode('utf-8')
    elif file_extension == '.pdf':
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(byte_object))
        return "".join(page.extract_text() for page in pdf_reader.pages)
    elif file_extension == '.docx':
        doc = Document(io.BytesIO(byte_object))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file format: {file_extension}")


async def _create_or_update_index(user: User, documents: List[LlamaDocument]) -> None:
    save_dir = f"./{user.username}_index"
    
    llm = OpenAI(model="gpt-3.5-turbo", temperature=0.1)
    embed_model = HuggingFaceEmbedding(model_name=user.embedder)

    service_context = ServiceContext.from_defaults(llm=llm, embed_model=embed_model)

    node_parser = HierarchicalNodeParser.from_defaults(
        chunk_sizes=[2048, 512, 128],
        chunk_overlap=20
    )
    nodes = node_parser.get_nodes_from_documents(documents)
    leaf_nodes = get_leaf_nodes(nodes)
    
    try:
        storage_context = StorageContext.from_defaults(persist_dir=save_dir)
        index = load_index_from_storage(storage_context, service_context=service_context)
        logger.info(f"Loaded existing index from {save_dir}")
        for node in nodes:
            storage_context.docstore.add_documents([node])
        for leaf_node in leaf_nodes:
            index.insert_nodes([leaf_node], show_progress=True)
    except (FileNotFoundError, ValueError):
        logger.info(f"Creating new index in {save_dir}")
        storage_context = StorageContext.from_defaults()
        storage_context.docstore.add_documents(nodes)
        index = VectorStoreIndex(leaf_nodes, storage_context=storage_context, service_context=service_context)
    
    index.storage_context.persist(persist_dir=save_dir)
    logger.info(f"Index persisted to {save_dir}")
@app.post("/question-answerer/")
async def ask_question(username: str = Form(...), question: str = Form(...), api_key: str = Form(...), confidence: float = Form(0.7)):
    os.environ["OPENAI_API_KEY"] = api_key
    index_dir = f"./{username}_index"
    if not os.path.exists(index_dir):
        logger.error(f"Index directory not found: {index_dir}")
        raise HTTPException(status_code=404, detail="No documents have been uploaded for this user.")

    temperature = 1.0 - confidence
    logger.info(f"Processing question for user: {username}")
    logger.info(f"Question: {question}")
    logger.info(f"Confidence: {confidence}, Temperature: {temperature}")

    try:
        llm = OpenAI(model="gpt-3.5-turbo", temperature=temperature, api_key=api_key)
        embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
        
        service_context = ServiceContext.from_defaults(llm=llm, embed_model=embed_model)
        storage_context = StorageContext.from_defaults(persist_dir=index_dir)
        
        logger.info("Loading index from storage")
        index = load_index_from_storage(storage_context, service_context=service_context)
        
        logger.info("Setting up AutoMergingRetriever")
        base_retriever = index.as_retriever(similarity_top_k=6)
        retriever = AutoMergingRetriever(
            base_retriever, 
            storage_context,
            verbose=True
        )
        
        logger.info("Creating RetrieverQueryEngine")
        query_engine = RetrieverQueryEngine.from_args(
            retriever,
            node_postprocessors=[],  
            verbose=True
        )

        logger.info("Querying the engine")
        response = query_engine.query(question)
        
        if response is None or response.response is None:
            logger.warning("Query engine returned None response")
            raise HTTPException(status_code=500, detail="No answer could be generated. Please try rephrasing your question.")
        
        answer = str(response)
        logger.info(f"Answer generated: {answer[:100]}...")  

        return {"answer": answer}
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An error occurred while processing your question: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error occurred: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")

@app.get("/check-index/{username}")
async def check_index(username: str):
    index_dir = f"./{username}_index"
    if not os.path.exists(index_dir):
        raise HTTPException(status_code=404, detail="Index not found for this user.")
    
    try:
        storage_context = StorageContext.from_defaults(persist_dir=index_dir)
        index = load_index_from_storage(storage_context)
        
        doc_count = len(storage_context.docstore.docs)
        node_count = len(index.index_struct.nodes)
        
        return {
            "document_count": doc_count,
            "node_count": node_count,
            "index_exists": True
        }
    except Exception as e:
        logger.error(f"Error checking index: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error checking index: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
